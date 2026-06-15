import os, datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0, 0, 0)
DARK = RGBColor(0x1a, 0x1a, 0x1a)
GREY = RGBColor(0x5a, 0x5a, 0x5a)
ACCENT = RGBColor(0x2b, 0x2b, 0x2b)
WHITE = RGBColor(0xff, 0xff, 0xff)
BANNER_BG = "1a1a1a"          # dark banner fill for the "bold" template
BANNER_CONTACT = RGBColor(0xcf, 0xca, 0xbf)

# ---------------------------------------------------------------------------
# Style system
# ---------------------------------------------------------------------------
# A "style" dict drives the look of the exported documents so the UI can offer
# a template gallery + accent / font / density controls. Everything is optional;
# an empty/None style reproduces the original editorial look exactly.
#
#   style = {
#     "template": "editorial" | "modern" | "classic" | "compact",
#     "accent":   "c8462e"  (hex, with or without leading '#'),
#     "font":     "Calibri" | "Georgia" | "Arial" | ...,
#     "density":  "comfortable" | "compact",
#   }

TEMPLATES = {
    "editorial": {
        "font": "Calibri",
        "name_align": WD_ALIGN_PARAGRAPH.CENTER,
        "name_size": 23,
        "name_color": BLACK,
        "heading_uses_accent": False,   # headings stay dark grey
        "heading_border": "999999",
        "heading_border_accent": False,
        "contact_align": WD_ALIGN_PARAGRAPH.CENTER,
    },
    "skyline": {
        "font": "Calibri",
        "name_align": WD_ALIGN_PARAGRAPH.LEFT,
        "name_size": 21,
        "name_color": BLACK,
        "heading_uses_accent": True,
        "heading_border": None,
        "heading_border_accent": True,
        "contact_align": WD_ALIGN_PARAGRAPH.LEFT,
        "default_accent": "3d8b7d",
    },
    "classic": {
        "font": "Georgia",
        "name_align": WD_ALIGN_PARAGRAPH.CENTER,
        "name_size": 22,
        "name_color": BLACK,
        "heading_uses_accent": False,
        "heading_color": BLACK,
        "heading_border": "000000",
        "heading_border_accent": False,
        "contact_align": WD_ALIGN_PARAGRAPH.CENTER,
    },
    "executive": {
        "font": "Calibri",
        "name_align": WD_ALIGN_PARAGRAPH.LEFT,
        "name_size": 18,
        "name_color": BLACK,
        "heading_uses_accent": True,
        "heading_border": None,
        "heading_border_accent": True,
        "contact_align": WD_ALIGN_PARAGRAPH.LEFT,
        "sidebar": True,
        "default_accent": "b8893f",
    },
    "aurora": {
        "font": "Calibri",
        "name_align": WD_ALIGN_PARAGRAPH.LEFT,
        "name_size": 18,
        "name_color": BLACK,
        "heading_uses_accent": True,
        "heading_border": None,
        "heading_border_accent": True,
        "contact_align": WD_ALIGN_PARAGRAPH.LEFT,
        "sidebar": True,
        "default_accent": "2f8f7d",
    },
    "serif": {
        "font": "Georgia",
        "name_align": WD_ALIGN_PARAGRAPH.CENTER,
        "name_size": 22,
        "name_color": BLACK,
        "heading_uses_accent": False,
        "heading_color": BLACK,
        "heading_border": "000000",
        "heading_border_accent": False,
        "contact_align": WD_ALIGN_PARAGRAPH.CENTER,
        "heading_align": WD_ALIGN_PARAGRAPH.CENTER,
        "heading_small_caps": True,
    },
    "spotlight": {
        "font": "Calibri",
        "name_align": WD_ALIGN_PARAGRAPH.LEFT,
        "name_size": 22,
        "name_color": WHITE,            # white on the coloured banner
        "heading_uses_accent": True,
        "heading_border": None,
        "heading_border_accent": True,
        "contact_align": WD_ALIGN_PARAGRAPH.LEFT,
        "banner": True,
        "banner_uses_accent": True,     # band filled with the accent colour
        "default_accent": "5f8d6e",
    },
    "minimal": {
        "font": "Calibri",
        "name_align": WD_ALIGN_PARAGRAPH.LEFT,
        "name_size": 19,
        "name_color": BLACK,
        "heading_uses_accent": False,
        "heading_color": RGBColor(0x22, 0x22, 0x22),
        "heading_border": "dddddd",     # hairline rule
        "heading_border_accent": False,
        "contact_align": WD_ALIGN_PARAGRAPH.LEFT,
    },
    "sidebar": {
        "font": "Calibri",
        "name_align": WD_ALIGN_PARAGRAPH.LEFT,
        "name_size": 18,
        "name_color": BLACK,
        "heading_uses_accent": True,
        "heading_border": None,
        "heading_border_accent": True,
        "contact_align": WD_ALIGN_PARAGRAPH.LEFT,
        "sidebar": True,                # two-column table layout
    },
}

SIDEBAR_BG = "f3efe6"                    # tinted left column (legacy)
SIDEBAR_RULE = "c9c4b8"                  # vertical divider rule for the sidebar layout

DEFAULT_TEMPLATE = "editorial"


def _rgb(hexstr, fallback=ACCENT):
    hexstr = (hexstr or "").lstrip("#")
    if len(hexstr) != 6:
        return fallback
    try:
        return RGBColor(int(hexstr[0:2], 16), int(hexstr[2:4], 16), int(hexstr[4:6], 16))
    except ValueError:
        return fallback


def _resolve_style(style):
    style = style or {}
    name = (style.get("template") or DEFAULT_TEMPLATE).lower()
    tmpl = TEMPLATES.get(name, TEMPLATES[DEFAULT_TEMPLATE])
    # Rich templates carry a curated accent; use it unless the user changed the swatch.
    picked = (style.get("accent") or "").lstrip("#")
    if tmpl.get("default_accent") and (not picked or picked == "c8462e"):
        accent = _rgb(tmpl["default_accent"], fallback=RGBColor(0xc8, 0x46, 0x2e))
    else:
        accent = _rgb(style.get("accent"), fallback=RGBColor(0xc8, 0x46, 0x2e))
    font = style.get("font") or tmpl["font"]
    density = (style.get("density") or "comfortable").lower()
    scale = 0.78 if density == "compact" else 1.0
    base_size = 9.5 if density == "compact" else 10.5

    name_color = tmpl.get("name_color")
    if tmpl.get("name_uses_accent"):
        name_color = accent
    if name_color is None:
        name_color = BLACK

    heading_color = accent if tmpl.get("heading_uses_accent") else tmpl.get("heading_color", ACCENT)
    heading_border = tmpl.get("heading_border")
    if tmpl.get("heading_border_accent"):
        heading_border = "%02x%02x%02x" % (accent[0], accent[1], accent[2])
    heading_border = heading_border or "999999"

    return {
        "template": name,
        "font": font,
        "accent": accent,
        "scale": scale,
        "base_size": base_size,
        "name_align": tmpl["name_align"],
        "name_size": tmpl["name_size"] * (0.85 if density == "compact" else 1.0),
        "name_color": name_color,
        "heading_color": heading_color,
        "heading_border": heading_border,
        "contact_align": tmpl["contact_align"],
        "heading_align": tmpl.get("heading_align", WD_ALIGN_PARAGRAPH.LEFT),
        "heading_small_caps": tmpl.get("heading_small_caps", False),
        "banner": tmpl.get("banner", False),
        "banner_bg": ("%02x%02x%02x" % (accent[0], accent[1], accent[2])
                      if tmpl.get("banner_uses_accent") else BANNER_BG),
        "sidebar": tmpl.get("sidebar", False),
    }


def _sp(p, before=0, after=0, line=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after)
    if line is not None: pf.line_spacing = line

def _border(p, size=6, color="000000", val="single"):
    el = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr"); b = OxmlElement("w:bottom")
    b.set(qn("w:val"), val); b.set(qn("w:sz"), str(size))
    b.set(qn("w:space"), "3"); b.set(qn("w:color"), color)
    pBdr.append(b); el.append(pBdr)

def _track(run, val):
    rPr = run._element.get_or_add_rPr(); s = OxmlElement("w:spacing")
    s.set(qn("w:val"), str(val)); rPr.append(s)

# Letter-spacing presets (twentieths of a point). Kept tasteful — wide tracking
# on a name reads as a rendering bug ("A M I T H"), not as design.
TRACK_NAME = 30
TRACK_HEADING = 24
TRACK_SMALLCAPS = 40

def _bullet(container, text, S, after=2.5, indent=0.22):
    """A hanging-indent bullet that renders identically in Word and LibreOffice.
    We draw the glyph ourselves (U+2022 in the body font) instead of relying on
    the 'List Bullet' style, whose Symbol-font bullet renders as tofu under some
    LibreOffice/font combinations."""
    p = container.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(indent)
    pf.first_line_indent = Inches(-indent)
    pf.tab_stops.add_tab_stop(Inches(indent))
    _sp(p, after=after, line=1.12)
    br = p.add_run("•\t"); br.font.size = Pt(S["base_size"]); br.font.color.rgb = S["accent"]
    tr = p.add_run(text); tr.font.size = Pt(S["base_size"]); tr.font.color.rgb = DARK
    return p

def _skill_groups(skills):
    """Normalise skills into [{'category': str|None, 'items': [str]}], accepting
    the grouped shape, a dict, or a legacy flat list of strings."""
    if not skills:
        return []
    if isinstance(skills, dict):
        return [{"category": k, "items": v} for k, v in skills.items() if v]
    if isinstance(skills, list) and skills and isinstance(skills[0], dict):
        out = []
        for g in skills:
            items = g.get("items") or []
            if items:
                out.append({"category": (g.get("category") or "").strip() or None, "items": items})
        return out
    return [{"category": None, "items": [s for s in skills if s]}]


def _doc_languages(resume):
    out = []
    for l in (resume.get("languages") or []):
        if isinstance(l, dict) and l.get("name"):
            out.append((l["name"], l.get("level") or ""))
        elif isinstance(l, str) and l.strip():
            out.append((l.strip(), ""))
    return out


def _entry(container, primary, topright, secondary, botright, bullets, S, tab_pos=6.7):
    """The ATS-standard entry layout:
        **Primary**                              Top-right (bold, e.g. location)
        Secondary (italic)                       Bottom-right (italic, e.g. dates)
        • bullets
    Used for experience (company/location, title/dates), projects and education."""
    head = container.add_paragraph(); _sp(head, before=7 * S["scale"], after=0, line=1.05)
    head.paragraph_format.tab_stops.add_tab_stop(Inches(tab_pos), WD_TAB_ALIGNMENT.RIGHT)
    r = head.add_run(primary or ""); r.bold = True
    r.font.size = Pt(S["base_size"] + 0.5); r.font.color.rgb = DARK
    if topright:
        tr = head.add_run(f"\t{topright}"); tr.bold = True
        tr.font.size = Pt(S["base_size"]); tr.font.color.rgb = DARK
    if secondary or botright:
        sub = container.add_paragraph(); _sp(sub, after=2, line=1.05)
        sub.paragraph_format.tab_stops.add_tab_stop(Inches(tab_pos), WD_TAB_ALIGNMENT.RIGHT)
        if secondary:
            sr = sub.add_run(secondary); sr.italic = True; sr.font.size = Pt(S["base_size"])
        if botright:
            br = sub.add_run(f"\t{botright}"); br.italic = True
            br.font.size = Pt(S["base_size"] - 1.5); br.font.color.rgb = GREY
    for b in (bullets or []):
        _bullet(container, b, S)

def _shade(cell, fill):
    """Fill a table cell with a solid background colour."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def _no_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + edge); e.set(qn("w:val"), "none"); e.set(qn("w:sz"), "0")
        borders.append(e)
    tblPr.append(borders)

def _cell_edge_border(cell, edge="right", color="cfcabf", size=8):
    """Draw a single border on one edge of a cell — used as the sidebar divider
    rule. A rule survives page breaks cleanly, unlike a shaded cell which leaves
    an empty colour band when the other column runs longer."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    e = OxmlElement("w:" + edge)
    e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(size))
    e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)
    borders.append(e); tcPr.append(borders)

def _banner(doc, name, contact, S, headline=""):
    """Full-width coloured banner holding the name + headline + contact."""
    table = doc.add_table(rows=1, cols=1)
    _no_table_borders(table)
    cell = table.cell(0, 0)
    _shade(cell, S.get("banner_bg", BANNER_BG))
    # tighten cell margins a touch via the single paragraph spacing
    np = cell.paragraphs[0]; _sp(np, before=2, after=2)
    nr = np.add_run((name or "").upper())
    nr.bold = True; nr.font.size = Pt(S["name_size"]); nr.font.color.rgb = WHITE; _track(nr, TRACK_NAME)
    if headline:
        hp = cell.add_paragraph(); _sp(hp, before=0, after=2)
        hr = hp.add_run(headline); hr.font.size = Pt(10); hr.font.color.rgb = WHITE; hr.bold = True
    if contact:
        cp = cell.add_paragraph(); _sp(cp, before=2, after=2)
        cr = cp.add_run(contact); cr.font.size = Pt(9); cr.font.color.rgb = BANNER_CONTACT
    doc.add_paragraph()  # spacer below the banner

def _set_cell_width(cell, width):
    cell.width = width
    tcPr = cell._tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:tcW")):
        tcPr.remove(existing)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width.twips))); tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)

def _cell_margins(cell, top=80, bottom=80, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right),
                      ("left", left), ("right", right)):
        m = OxmlElement("w:" + edge); m.set(qn("w:w"), str(val)); m.set(qn("w:type"), "dxa")
        mar.append(m)
    tcPr.append(mar)

def _cell_heading(cell, text, S, first=False):
    p = cell.add_paragraph(); _sp(p, before=0 if first else 10, after=3)
    r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(S["base_size"])
    r.font.color.rgb = S["heading_color"]; _track(r, TRACK_HEADING)
    return p

def _job_into(container, job, S, name_key="title", link_key="company"):
    # Sidebar/right-cell experience: company (bold) / title (italic) + dates.
    _entry(container, job.get(link_key) or job.get(name_key, ""),
           job.get("location"),
           job.get(name_key) if job.get(link_key) else None,
           job.get("dates"), job.get("bullets") or [], S, tab_pos=4.0)

def _build_resume_sidebar(resume, doc, S):
    """Two-column layout: tinted left column (contact, skills, education),
    main right column (summary, experience, projects). Rendered as a borderless
    table so it survives the .docx/PDF pipeline."""
    table = doc.add_table(rows=1, cols=2)
    table.allow_autofit = False
    _no_table_borders(table)
    left, right = table.cell(0, 0), table.cell(0, 1)
    _set_cell_width(left, Inches(2.3)); _set_cell_width(right, Inches(4.5))
    # A vertical divider rule (not cell shading) gives the sidebar feel while
    # surviving page breaks cleanly — no empty colour band on overflow pages.
    _cell_edge_border(left, "right", color=SIDEBAR_RULE)
    _cell_margins(left, left=20, right=180); _cell_margins(right, left=200, right=20)

    c = resume.get("contact", {}) or {}

    # ---- left column ----
    np = left.paragraphs[0]; _sp(np, after=2 if (resume.get("headline") or "").strip() else 5)
    nr = np.add_run((resume.get("name", "") or "").upper())
    nr.bold = True; nr.font.size = Pt(S["name_size"]); nr.font.color.rgb = BLACK; _track(nr, TRACK_NAME)
    headline = (resume.get("headline") or "").strip()
    if headline:
        hp = left.add_paragraph(); _sp(hp, after=5)
        hr = hp.add_run(headline); hr.font.size = Pt(9); hr.bold = True; hr.font.color.rgb = S["heading_color"]
    for v in [c.get("email"), c.get("phone"), c.get("location"), *(c.get("links") or [])]:
        if v:
            cp = left.add_paragraph(); _sp(cp, after=1, line=1.05)
            cr = cp.add_run(v); cr.font.size = Pt(8.5); cr.font.color.rgb = GREY

    if resume.get("skills"):
        _cell_heading(left, "Skills", S)
        for g in _skill_groups(resume["skills"]):
            sp = left.add_paragraph(); _sp(sp, after=3, line=1.2)
            if g["category"]:
                cr = sp.add_run(f"{g['category']}\n"); cr.bold = True; cr.font.size = Pt(S["base_size"] - 0.5)
            ir = sp.add_run(", ".join(g["items"])); ir.font.size = Pt(S["base_size"] - 0.5)

    if resume.get("education"):
        _cell_heading(left, "Education", S)
        for e in resume["education"]:
            ep = left.add_paragraph(); _sp(ep, after=3)
            ep.add_run(e.get("degree", "")).bold = True
            if e.get("institution"):
                ir = ep.add_run(f"\n{e['institution']}"); ir.font.size = Pt(9); ir.font.color.rgb = GREY
            if e.get("dates"):
                dr = ep.add_run(f"\n{e['dates']}"); dr.italic = True; dr.font.size = Pt(8.5); dr.font.color.rgb = GREY

    if resume.get("certifications"):
        _cell_heading(left, "Certifications", S)
        for cert in resume["certifications"]:
            cp = left.add_paragraph(); _sp(cp, after=2)
            cp.add_run(cert).font.size = Pt(9)

    langs = _doc_languages(resume)
    if langs:
        _cell_heading(left, "Languages", S)
        for nm, lv in langs:
            lp = left.add_paragraph(); _sp(lp, after=1)
            lr = lp.add_run(nm); lr.font.size = Pt(9)
            if lv:
                vr = lp.add_run(f" — {lv}"); vr.font.size = Pt(8.5); vr.font.color.rgb = GREY

    # ---- right column ----
    first = True
    if resume.get("summary"):
        _cell_heading(right, "Professional Summary", S, first=True); first = False
        p = right.add_paragraph(resume["summary"]); _sp(p, after=2, line=1.12)
    if resume.get("experience"):
        _cell_heading(right, "Professional Experience", S, first=first); first = False
        for job in resume["experience"]:
            _job_into(right, job, S, "title", "company")
    if resume.get("projects"):
        _cell_heading(right, "Projects", S, first=first)
        for proj in resume["projects"]:
            _job_into(right, proj, S, "name", "link")

def _base(S):
    doc = Document(); s = doc.sections[0]
    s.page_width = Inches(8.5); s.page_height = Inches(11)
    s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.6)
    s.left_margin = Inches(0.85); s.right_margin = Inches(0.85)
    n = doc.styles["Normal"]; n.font.name = S["font"]
    n.font.size = Pt(S["base_size"]); n.font.color.rgb = DARK
    n.paragraph_format.space_after = Pt(0); n.paragraph_format.line_spacing = 1.1
    return doc

def _heading(doc, text, S):
    p = doc.add_paragraph(); _sp(p, before=12 * S["scale"], after=5 * S["scale"])
    p.alignment = S.get("heading_align", WD_ALIGN_PARAGRAPH.LEFT)
    small_caps = S.get("heading_small_caps", False)
    # Small caps needs the original case; otherwise we upper-case for the rule look.
    r = p.add_run(text if small_caps else text.upper())
    r.bold = True; r.font.size = Pt(S["base_size"]); r.font.color.rgb = S["heading_color"]
    if small_caps:
        r.font.small_caps = True
    _track(r, TRACK_SMALLCAPS if small_caps else TRACK_HEADING)
    _border(p, size=4, color=S["heading_border"])
    return p

def build_resume(resume, out_path, style=None):
    S = _resolve_style(style)
    doc = _base(S)

    if S.get("sidebar"):
        _build_resume_sidebar(resume, doc, S)
        doc.save(out_path); return out_path

    c = resume.get("contact", {}) or {}
    bits = [c.get("email"), c.get("phone"), c.get("location"), *(c.get("links") or [])]
    contact = "    •    ".join(b for b in bits if b)

    headline = (resume.get("headline") or "").strip()
    if S.get("banner"):
        _banner(doc, resume.get("name", ""), contact, S, headline=headline)
    else:
        name_p = doc.add_paragraph(); name_p.alignment = S["name_align"]
        _sp(name_p, after=1 if headline else 3)
        nr = name_p.add_run((resume.get("name","") or "").upper())
        nr.bold = True; nr.font.size = Pt(S["name_size"]); nr.font.color.rgb = S["name_color"]; _track(nr, TRACK_NAME)
        if headline:
            hp = doc.add_paragraph(); hp.alignment = S["name_align"]; _sp(hp, after=4)
            hr = hp.add_run(headline); hr.bold = True; hr.font.size = Pt(10.5); hr.font.color.rgb = S["heading_color"]
        if contact:
            cp = doc.add_paragraph(); cp.alignment = S["contact_align"]
            _sp(cp, after=6); cr = cp.add_run(contact)
            cr.font.size = Pt(9); cr.font.color.rgb = GREY
            _border(cp, size=8, color="000000")

    if resume.get("summary"):
        _heading(doc, "Professional Summary", S)
        p = doc.add_paragraph(resume["summary"]); _sp(p, after=2, line=1.12)

    if resume.get("skills"):
        _heading(doc, "Skills", S)
        for g in _skill_groups(resume["skills"]):
            p = doc.add_paragraph(); _sp(p, after=2, line=1.18)
            if g["category"]:
                cr = p.add_run(f"{g['category']}: "); cr.bold = True; cr.font.size = Pt(S["base_size"])
            ir = p.add_run(", ".join(g["items"])); ir.font.size = Pt(S["base_size"])

    if resume.get("experience"):
        _heading(doc, "Experience", S)
        for job in resume["experience"]:
            _entry(doc, job.get("company") or job.get("title", ""), job.get("location"),
                   job.get("title") if job.get("company") else None, job.get("dates"),
                   job.get("bullets") or [], S)

    if resume.get("projects"):
        _heading(doc, "Projects", S)
        for proj in resume["projects"]:
            _entry(doc, proj.get("name", ""), proj.get("dates"), proj.get("link"),
                   None, proj.get("bullets") or [], S)

    if resume.get("education"):
        _heading(doc, "Education", S)
        for ed in resume["education"]:
            _entry(doc, ed.get("institution") or ed.get("degree", ""), ed.get("dates"),
                   ed.get("degree") if ed.get("institution") else None, None, [], S)

    if resume.get("certifications"):
        _heading(doc, "Certifications", S)
        for cert in resume["certifications"]:
            _bullet(doc, cert, S)

    langs = _doc_languages(resume)
    if langs:
        _heading(doc, "Languages", S)
        p = doc.add_paragraph(); _sp(p, after=2, line=1.2)
        for i, (nm, lv) in enumerate(langs):
            if i:
                p.add_run("     ").font.size = Pt(S["base_size"])
            r1 = p.add_run(nm); r1.font.size = Pt(S["base_size"])
            if lv:
                r2 = p.add_run(f"  {lv}"); r2.bold = True; r2.font.size = Pt(S["base_size"])
                r2.font.color.rgb = S["heading_color"]
    doc.save(out_path); return out_path

def build_cover_letter(letter, applicant_name, out_path, contact_line="", links=None, style=None):
    S = _resolve_style(style)
    doc = _base(S)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.2
    if applicant_name:
        n = doc.add_paragraph(); n.alignment = S["name_align"]; _sp(n, after=2)
        nr = n.add_run(applicant_name.upper()); nr.bold = True
        nr.font.size = Pt(max(16, S["name_size"] - 4)); nr.font.color.rgb = S["name_color"]; _track(nr, TRACK_NAME)
    line_bits = []
    if contact_line: line_bits.append(contact_line)
    if links: line_bits.extend(links)
    if line_bits:
        cp = doc.add_paragraph(); cp.alignment = S["contact_align"]; _sp(cp, after=6)
        cr = cp.add_run("   •   ".join(line_bits)); cr.font.size = Pt(9)
        cr.font.color.rgb = GREY; _border(cp, size=8, color="000000")
    dp = doc.add_paragraph(); _sp(dp, before=2, after=10)
    dp.add_run(datetime.date.today().strftime("%B %d, %Y")).font.color.rgb = GREY

    g = doc.add_paragraph(letter.get("greeting","Dear Hiring Manager,")); _sp(g, after=9)
    for para in letter.get("body_paragraphs", []):
        bp = doc.add_paragraph(para); _sp(bp, after=9, line=1.2)
    cl = doc.add_paragraph(letter.get("closing","Sincerely,")); _sp(cl, before=4, after=1)
    if applicant_name: doc.add_paragraph(applicant_name)
    doc.save(out_path); return out_path
import shutil, subprocess

def to_pdf(docx_path):
    """Convert a .docx to .pdf via LibreOffice. Returns the pdf path, or None
    if LibreOffice isn't available / conversion fails (caller falls back to docx)."""
    return to_pdfs([docx_path]).get(docx_path)


def to_pdfs(docx_paths):
    """Convert multiple .docx files in one LibreOffice run.

    Starting LibreOffice is the slow part, so batching resume + cover letter
    avoids paying that startup cost twice.
    """
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        mac = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        soffice = mac if os.path.exists(mac) else None
    if not soffice:
        return {p: None for p in docx_paths}
    if not docx_paths:
        return {}
    outdir = os.path.dirname(docx_paths[0])
    try:
        subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", outdir, *docx_paths],
            check=True, capture_output=True, timeout=60)
    except Exception:
        return {p: None for p in docx_paths}
    out = {}
    for docx_path in docx_paths:
        pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
        out[docx_path] = pdf_path if os.path.exists(pdf_path) else None
    return out
