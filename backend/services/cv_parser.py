from __future__ import annotations
"""
Extract plain text from an uploaded CV.

Supports .pdf and .docx. Returns (text, ok). If extraction yields too little
usable text — scanned/image-only PDF, corrupt file, password protection — ok is
False and the caller prompts the user to paste their CV instead.
"""
import os
import pdfplumber
from docx import Document

# Below this many characters we assume extraction effectively failed
# (e.g. an image-only scanned PDF where pdfplumber finds no text layer).
MIN_USABLE_CHARS = 120

# pdfplumber groups words into lines using y_tolerance (default 3pt). Stylised
# headers — e.g. a name whose first capital of each word is set a few points
# larger/higher than the rest — get split into bogus lines ("A A R" / "MITH
# AJOLKAR" instead of "AMITH A RAJOLKAR"). A slightly larger tolerance merges
# those raised caps back onto their real line while staying well below normal
# body line spacing (typically >=10pt), so paragraphs are unaffected.
PDF_Y_TOLERANCE = 6


def _from_pdf(path: str) -> str:
    chunks = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            txt = page.extract_text(y_tolerance=PDF_Y_TOLERANCE) or ""
            if txt.strip():
                chunks.append(txt)
    return "\n".join(chunks).strip()


def _from_docx(path: str) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # Pull text out of tables too — many CVs lay out contact info in tables.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts).strip()


def parse_cv(path: str) -> tuple[str, bool]:
    """Return (extracted_text, ok)."""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            text = _from_pdf(path)
        elif ext in (".docx", ".doc"):
            # python-docx cannot read legacy .doc; those will raise and fall back.
            text = _from_docx(path)
        else:
            return "", False
    except Exception:
        return "", False

    return (text, len(text) >= MIN_USABLE_CHARS)


def looks_like_cv(text: str) -> bool:
    """Light sanity check on pasted text so we don't tailor against gibberish."""
    return bool(text) and len(text.strip()) >= MIN_USABLE_CHARS
